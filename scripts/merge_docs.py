"""Merge all .docx files from Research/Merged into a single document."""

from docx import Document
from docx.shared import Pt, RGBColor
from pathlib import Path
import os
import sys

# Force UTF-8 encoding for console output
if sys.stdout.encoding != 'utf-8':
    sys.stdout.reconfigure(encoding='utf-8', errors='replace')

def merge_docx_files(source_dir, output_file):
    """Merge all .docx files in source_dir into output_file."""

    # Get all docx files, sorted alphabetically
    source_path = Path(source_dir)
    docx_files = sorted(source_path.glob('*.docx'))

    print(f"Found {len(docx_files)} documents to merge")

    # Create new combined document
    combined = Document()

    # Add title page
    title = combined.add_heading('DnD 3.5e AIDM Research Documents', 0)
    combined.add_paragraph(f'Combined from {len(docx_files)} source documents')
    combined.add_paragraph(f'Generated: {Path.cwd()}')
    combined.add_paragraph('')

    # Add table of contents
    combined.add_heading('Table of Contents', level=1)
    for i, docx_file in enumerate(docx_files, 1):
        combined.add_paragraph(f'{i}. {docx_file.name}', style='List Number')

    # Merge each document
    for i, docx_file in enumerate(docx_files, 1):
        print(f'Processing {i}/{len(docx_files)}: {docx_file.name}')

        # Add page break and document header
        combined.add_page_break()
        combined.add_heading(f'[{i}] {docx_file.stem}', level=1)
        combined.add_paragraph(f'Source: {docx_file.name}')
        combined.add_paragraph('')

        # Read source document
        try:
            source = Document(docx_file)

            # Copy all elements from source document
            for para in source.paragraphs:
                # Create new paragraph with same text
                new_para = combined.add_paragraph(para.text, style=para.style)

                # Copy paragraph formatting
                if para.style:
                    try:
                        new_para.style = para.style
                    except:
                        pass  # Style might not exist in target

            # Copy tables
            for table in source.tables:
                # Create new table with same dimensions
                new_table = combined.add_table(rows=len(table.rows), cols=len(table.columns))

                # Copy cell contents
                for i, row in enumerate(table.rows):
                    for j, cell in enumerate(row.cells):
                        new_table.rows[i].cells[j].text = cell.text

            print(f'  [OK] Merged: {len(source.paragraphs)} paragraphs, {len(source.tables)} tables')

        except Exception as e:
            print(f'  [WARNING] Error processing {docx_file.name}: {e}')
            combined.add_paragraph(f'[ERROR: Could not merge this document - {e}]')

    # Save combined document
    output_path = Path(output_file)
    combined.save(output_path)
    print(f'\n[SUCCESS] Combined document saved to: {output_path}')
    print(f'   Total size: {output_path.stat().st_size / 1024:.1f} KB')

if __name__ == '__main__':
    source_dir = r'F:\DnD-3.5\Research\Merged'
    output_file = r'F:\DnD-3.5\Research\Combined_Research_Documents.docx'

    merge_docx_files(source_dir, output_file)
